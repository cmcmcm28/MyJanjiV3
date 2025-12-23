# Contract Service for PDF Generation
# Handles: template download, placeholder filling, PDF conversion, upload

import os
import re
import json
import tempfile
import io
import base64
import requests
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from supabase import create_client, Client
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY")  # Use service key for storage access

# Storage bucket names
TEMPLATE_BUCKET = "contract_templates"   # Bucket for .docx templates
PDF_BUCKET = "contract_pdf"              # Bucket for generated PDFs
GENERATED_FOLDER = "generated"

# Template ID to file path mapping (loaded from config.json or hardcoded fallback)
TEMPLATE_PATHS = {
    "VEHICLE_USE": "Items & Assets/Vehicle Borrowing.docx",
    "ITEM_BORROW": "Items & Assets/Item Borrowing.docx",
    "BILL_SPLIT": "Money & Finance/Bill Split.docx",
    "FRIENDLY_LOAN": "Money & Finance/Friendly Loan.docx",
    "FREELANCE_JOB": "Service & Gig Work/Freelance Job.docx",
    "SALE_DEPOSIT": "Service & Gig Work/Sales & Deposit.docx",
}

# Local temp folder
TEMP_FOLDER = "temp_contracts"
os.makedirs(TEMP_FOLDER, exist_ok=True)


def get_supabase_client() -> Client:
    """Get Supabase client instance"""
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise ValueError("SUPABASE_URL and SUPABASE_SERVICE_KEY must be set in .env")
    return create_client(SUPABASE_URL, SUPABASE_KEY)


def get_template_path(template_id: str) -> str:
    """
    Get the storage path for a template ID.
    Uses TEMPLATE_PATHS mapping or treats input as direct path.
    """
    # If it's a known template ID, map it
    if template_id in TEMPLATE_PATHS:
        return TEMPLATE_PATHS[template_id]
    
    # If it already looks like a path (contains / or .docx), use as-is
    if '/' in template_id or template_id.endswith('.docx'):
        return template_id
    
    # Default: assume it's in root with .docx extension
    return f"{template_id}.docx"


def download_template(template_name: str) -> str:
    """
    Download .docx template from Supabase Storage.
    Accepts template ID (e.g., 'ITEM_BORROW') or full path.
    Returns local file path.
    """
    supabase = get_supabase_client()
    
    # Resolve template ID to storage path
    storage_path = get_template_path(template_name)
    
    print(f"📥 Downloading template: {storage_path}")
    
    try:
        # Try using signed URL (handles special characters better)
        signed_url_response = supabase.storage.from_(TEMPLATE_BUCKET).create_signed_url(storage_path, 60)
        if signed_url_response and 'signedURL' in signed_url_response:
            url = signed_url_response['signedURL']
            print(f"📡 Using signed URL")
            
            response = requests.get(url)
            if response.status_code == 200:
                # Save to local temp file
                filename = os.path.basename(storage_path)
                local_path = os.path.join(TEMP_FOLDER, filename)
                with open(local_path, 'wb') as f:
                    f.write(response.content)
                
                print(f"✅ Template saved to: {local_path}")
                return local_path
    except Exception as e:
        print(f"⚠️ Signed URL failed: {e}")
    
    # Fallback: try direct download
    try:
        print(f"📡 Trying direct download...")
        response = supabase.storage.from_(TEMPLATE_BUCKET).download(storage_path)
        
        # Save to local temp file
        filename = os.path.basename(storage_path)
        local_path = os.path.join(TEMP_FOLDER, filename)
        with open(local_path, 'wb') as f:
            f.write(response)
        
        print(f"✅ Template saved to: {local_path}")
        return local_path
    except Exception as e:
        raise Exception(f"Failed to download template '{storage_path}': {e}")


def fetch_image(image_source):
    """
    Download image from URL or decode base64.
    Returns bytes stream (io.BytesIO) or None.
    """
    try:
        if not image_source:
            return None
            
        # Case 1: Base64 string
        if isinstance(image_source, str) and image_source.startswith('data:image'):
            # format: "data:image/png;base64,iVBQR..."
            header, encoded = image_source.split(',', 1)
            return io.BytesIO(base64.b64decode(encoded))
            
        # Case 2: URL
        if isinstance(image_source, str) and (image_source.startswith('http://') or image_source.startswith('https://')):
            response = requests.get(image_source, timeout=10)
            if response.status_code == 200:
                return io.BytesIO(response.content)
                
        return None
    except Exception as e:
        print(f"⚠️ Failed to fetch image: {e}")
        return None


def fill_template(doc_path: str, placeholders: dict) -> str:
    """
    Replace {{PLACEHOLDER}} with actual values in the .docx document.
    Handles text replacement and image insertion for signatures.
    Returns path to the filled document.
    """
    print(f"📝 Filling template with {len(placeholders)} placeholders")
    
    doc = Document(doc_path)
    
    # Identify signature keys that should be treated as images
    signature_keys = ['CREATOR_SIGNATURE', 'ACCEPTEE_SIGNATURE', 'SIGNATURE']
    
    def process_paragraph(paragraph):
        for key, value in placeholders.items():
            placeholder = f"{{{{{key}}}}}"  # {{KEY}}
            
            if placeholder in paragraph.text:
                # Special handling for signatures (images)
                if key in signature_keys and value:
                    print(f"🖼️ Found signature placeholder: {key}")
                    # Clear the placeholder text
                    # We need to find the specific run containing the placeholder
                    # Note: This simple replacement assumes placeholder is not split across runs
                    # For a robust solution, we iterate runs.
                    
                    found = False
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            # Remove the placeholder text
                            run.text = run.text.replace(placeholder, "")
                            
                            # Insert image
                            img_stream = fetch_image(value)
                            if img_stream:
                                run.add_picture(img_stream, width=Inches(1.5))
                                print(f"✅ Signature inserted for {key}")
                                found = True
                    
                else:
                    # Text replacement
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    
    # Save filled document
    filled_path = doc_path.replace('.docx', '_filled.docx')
    doc.save(filled_path)
    
    print(f"✅ Filled template saved to: {filled_path}")
    return filled_path


def convert_to_pdf(docx_path: str) -> str:
    """
    Convert .docx to PDF using Microsoft Word.
    Returns path to the PDF file.
    """
    print(f"🔄 Converting to PDF...")
    
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        convert(docx_path, pdf_path)
        print(f"✅ PDF created: {pdf_path}")
        return pdf_path
    except AttributeError as e:
        # Known issue: Word.Application.Quit fails but PDF is created successfully
        if "Quit" in str(e) and os.path.exists(pdf_path):
            print(f"⚠️ Word quit error ignored, PDF created: {pdf_path}")
            return pdf_path
        print(f"❌ PDF conversion failed: {e}")
        raise
    except Exception as e:
        # Check if PDF was created despite error
        if os.path.exists(pdf_path):
            print(f"⚠️ Error occurred but PDF exists: {pdf_path}")
            return pdf_path
        print(f"❌ PDF conversion failed: {e}")
        raise


def upload_pdf(pdf_path: str, contract_id: str) -> str:
    """
    Upload generated PDF to Supabase Storage (contract_pdf bucket).
    Returns public URL.
    """
    supabase = get_supabase_client()
    
    # Storage path in the PDF bucket
    storage_path = f"{GENERATED_FOLDER}/{contract_id}.pdf"
    
    print(f"📤 Uploading PDF to bucket '{PDF_BUCKET}': {storage_path}")
    
    # Read PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()
    
    # Upload to Supabase PDF bucket
    response = supabase.storage.from_(PDF_BUCKET).upload(
        storage_path,
        pdf_data,
        file_options={"content-type": "application/pdf"}
    )
    
    # Get public URL
    public_url = supabase.storage.from_(PDF_BUCKET).get_public_url(storage_path)
    
    print(f"✅ PDF uploaded: {public_url}")
    return public_url


def generate_contract(template_name: str, placeholders: dict, contract_id: str) -> dict:
    """
    Full workflow: download template → fill placeholders → convert to PDF → upload.
    Returns dict with success status and PDF URL.
    """
    try:
        # Step 1: Download template
        doc_path = download_template(template_name)
        
        # Step 2: Fill placeholders
        filled_path = fill_template(doc_path, placeholders)
        
        # Step 3: Convert to PDF
        pdf_path = convert_to_pdf(filled_path)
        
        # Step 4: Upload PDF
        pdf_url = upload_pdf(pdf_path, contract_id)
        
        # Cleanup temp files
        cleanup_temp_files([doc_path, filled_path, pdf_path])
        
        return {
            "success": True,
            "pdf_url": pdf_url,
            "contract_id": contract_id
        }
        
    except Exception as e:
        print(f"❌ Contract generation failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }


def cleanup_temp_files(file_paths: list):
    """Remove temporary files"""
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass


# Preview function (returns filled docx without converting to PDF)
def preview_contract(template_name: str, placeholders: dict) -> str:
    """
    Generate preview of contract (filled .docx).
    Returns local path to filled document.
    """
    doc_path = download_template(template_name)
    filled_path = fill_template(doc_path, placeholders)
    return filled_path
