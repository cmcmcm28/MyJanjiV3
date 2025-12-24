# Contract Service for PDF Generation
# Handles: template download, placeholder filling, PDF conversion, upload

import os
import re
import json
import tempfile
import io
import base64
import requests
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from docx2pdf import convert
from supabase import create_client, Client

from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Supabase configuration
SUPABASE_URL = os.getenv("SUPABASE_URL")
# Use service key for storage access
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

# Storage bucket names
TEMPLATE_BUCKET = "contract_templates"   # Bucket for .docx templates
PDF_BUCKET = "contract-pdf"              # Bucket for generated PDFs
GENERATED_FOLDER = "generated"

# ============================================
# TEMPLATE CACHE - Store downloaded templates in memory
# ============================================
TEMPLATE_CACHE = {}  # { "storage_path": (file_bytes, timestamp) }
CACHE_EXPIRY_SECONDS = 3600  # 1 hour cache expiry


def get_cached_template(storage_path: str):
    """
    Get template from cache if available and not expired.
    Returns (bytes, True) if cache hit, (None, False) if cache miss.
    """
    if storage_path in TEMPLATE_CACHE:
        cached_bytes, cached_time = TEMPLATE_CACHE[storage_path]
        age = time.time() - cached_time
        if age < CACHE_EXPIRY_SECONDS:
            print(f"✅ Cache HIT for '{storage_path}' (age: {age:.1f}s)")
            return cached_bytes, True
        else:
            print(f"⏰ Cache EXPIRED for '{storage_path}' (age: {age:.1f}s)")
            del TEMPLATE_CACHE[storage_path]
    return None, False


def set_cached_template(storage_path: str, file_bytes: bytes):
    """Store template bytes in cache with current timestamp."""
    TEMPLATE_CACHE[storage_path] = (file_bytes, time.time())
    print(f"💾 Cached template: '{storage_path}' ({len(file_bytes)} bytes)")


def clear_template_cache():
    """Clear all cached templates (useful for admin/debug)."""
    TEMPLATE_CACHE.clear()
    print("🗑️ Template cache cleared")


# ============================================

# Template ID to file path mapping (loaded from config.json or hardcoded fallback)
TEMPLATE_PATHS = {
    "VEHICLE_USE": "Items & Assets/Vehicle Borrowing.docx",
    "ITEM_BORROW": "Items & Assets/Item Borrowing.docx",
    "BILL_SPLIT": "Money & Finance/Bill Split.docx",
    "FRIENDLY_LOAN": "Money & Finance/Friendly Loan.docx",
    "FREELANCE_JOB": "Service & Gig Work/Freelance Job.docx",
    "SALE_DEPOSIT": "Service & Gig Work/Sales & Deposit.docx",
}

# Load template config for mapping
try:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    CONFIG_PATH = os.path.join(BASE_DIR, 'templates_config.json')
    with open(CONFIG_PATH, 'r') as f:
        TEMPLATE_CONFIG = json.load(f)
except Exception as e:
    print(f"Warning: Failed to load templates_config.json: {e}")
    TEMPLATE_CONFIG = {"categories": []}


def get_template_mapping(template_id: str) -> dict:
    """Find mapping for a given template ID"""
    for category in TEMPLATE_CONFIG.get('categories', []):
        for template in category.get('templates', []):
            if template['id'] == template_id:
                return template.get('mapping', {})
    return {}


def map_placeholders(placeholders: dict, mapping: dict) -> dict:
    """
    Map frontend keys to Docx placeholders based on config.
    If a key is not in mapping, it is passed through as-is (fallback).
    """
    mapped_data = {}

    # 1. Apply mapping
    for frontend_key, docx_key in mapping.items():
        if frontend_key in placeholders:
            mapped_data[docx_key] = placeholders[frontend_key]

    # 2. Pass through any keys that didn't match mapping (in case frontend sends raw keys)
    for key, value in placeholders.items():
        if key not in mapping:
            mapped_data[key] = value

    return mapped_data


# Local temp folder for working files
TEMP_FOLDER = "temp_contracts"
os.makedirs(TEMP_FOLDER, exist_ok=True)

# Prepared contracts folder (stores pre-generated PDFs awaiting preview/signing)
PREPARED_FOLDER = "prepared_contracts"
os.makedirs(PREPARED_FOLDER, exist_ok=True)

# Signatures folder (stores temporary signature images)
SIGNATURES_FOLDER = "temp_signatures"
os.makedirs(SIGNATURES_FOLDER, exist_ok=True)


def get_supabase_client() -> Client:
    """Get Supabase client instance"""
    if not SUPABASE_URL or not SUPABASE_KEY:
        raise ValueError(
            "SUPABASE_URL and SUPABASE_SERVICE_KEY must be set in .env")
    return create_client(SUPABASE_URL, SUPABASE_KEY)


def get_template_path(template_id: str) -> str:
    """
    Get the storage path for a template ID.
    Uses TEMPLATE_PATHS mapping or treats input as direct path.
    """
    # If it's a known template ID, map it
    if template_id in TEMPLATE_PATHS:
        return TEMPLATE_PATHS[template_id]

    # If it already looks like a path (contains / or .docx), use as-is
    if '/' in template_id or template_id.endswith('.docx'):
        return template_id

    # Default: assume it's in root with .docx extension
    return f"{template_id}.docx"


def download_template(template_name: str) -> str:
    """
    Download .docx template from Supabase Storage (with caching).
    Accepts template ID (e.g., 'ITEM_BORROW') or full path.
    Returns local file path.

    Caching: Templates are cached in memory for 1 hour to avoid
    repeated downloads of the same template file.
    """
    # Resolve template ID to storage path
    storage_path = get_template_path(template_name)

    # Check cache first
    cached_bytes, cache_hit = get_cached_template(storage_path)

    if cache_hit:
        # Use cached template - save to local file and return
        filename = os.path.basename(storage_path)
        local_path = os.path.join(TEMP_FOLDER, filename)
        with open(local_path, 'wb') as f:
            f.write(cached_bytes)
        return local_path

    # Cache miss - download from Supabase
    supabase = get_supabase_client()
    print(f"📥 Downloading template: {storage_path}")

    file_bytes = None

    try:
        # Try using signed URL (handles special characters better)
        signed_url_response = supabase.storage.from_(
            TEMPLATE_BUCKET).create_signed_url(storage_path, 60)
        if signed_url_response and 'signedURL' in signed_url_response:
            url = signed_url_response['signedURL']
            print(f"Using signed URL")

            response = requests.get(url)
            if response.status_code == 200:
                file_bytes = response.content
    except Exception as e:
        print(f"Signed URL failed: {e}")

    # Fallback: try direct download
    if file_bytes is None:
        try:
            print(f"Trying direct download...")
            file_bytes = supabase.storage.from_(
                TEMPLATE_BUCKET).download(storage_path)
        except Exception as e:
            raise Exception(
                f"Failed to download template '{storage_path}': {e}")

    if file_bytes is None:
        raise Exception(f"Failed to download template '{storage_path}'")

    # Cache the downloaded template
    set_cached_template(storage_path, file_bytes)

    # Save to local temp file
    filename = os.path.basename(storage_path)
    local_path = os.path.join(TEMP_FOLDER, filename)
    with open(local_path, 'wb') as f:
        f.write(file_bytes)

    print(f"Template saved to: {local_path}")
    return local_path


def fetch_image(image_source):
    """
    Download image from URL or decode base64.
    Returns bytes stream (io.BytesIO) or None.
    """
    try:
        if not image_source:
            return None

        # Case 1: Base64 string
        if isinstance(image_source, str) and image_source.startswith('data:image'):
            # format: "data:image/png;base64,iVBQR..."
            header, encoded = image_source.split(',', 1)
            return io.BytesIO(base64.b64decode(encoded))

        # Case 2: URL
        if isinstance(image_source, str) and (image_source.startswith('http://') or image_source.startswith('https://')):
            response = requests.get(image_source, timeout=10)
            if response.status_code == 200:
                return io.BytesIO(response.content)

        return None
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        return None


def fill_template(doc_path: str, placeholders: dict) -> str:
    """
    Replace {{PLACEHOLDER}} with actual values in the .docx document.
    Handles text replacement and image insertion for signatures.
    Returns path to the filled document.
    """
    print(f"Filling template with {len(placeholders)} placeholders")

    doc = Document(doc_path)

    # Identify signature keys that should be treated as images (both upper and lowercase)
    signature_keys = ['CREATOR_SIGNATURE', 'ACCEPTEE_SIGNATURE', 'ACCEPTOR_SIGNATURE', 'SIGNATURE',
                      'creator_signature', 'acceptee_signature', 'acceptor_signature', 'signature']

    # Keys that should be formatted with bold and yellow highlight (like creator name fields)
    highlighted_keys = [
        # Body acceptee fields
        'ACCEPTEE_NAME', 'acceptee_name', 'ACCEPTOR_NAME', 'acceptor_name',
        'ACCEPTEE_IC', 'acceptee_ic', 'ACCEPTOR_IC', 'acceptor_ic',
        'acceptee_id_number', 'acceptor_id_number',
        # Signature section fields
        'creator_signature_name', 'CREATOR_SIGNATURE_NAME',
        'creator_signature_id', 'CREATOR_SIGNATURE_ID',
        'creator_signature_date', 'CREATOR_SIGNATURE_DATE',
        'acceptee_signature_name', 'ACCEPTEE_SIGNATURE_NAME', 'acceptor_signature_name', 'ACCEPTOR_SIGNATURE_NAME',
        'acceptee_signature_id', 'ACCEPTEE_SIGNATURE_ID', 'acceptor_signature_id', 'ACCEPTOR_SIGNATURE_ID',
        'acceptee_signature_date', 'ACCEPTEE_SIGNATURE_DATE', 'acceptor_signature_date', 'ACCEPTOR_SIGNATURE_DATE',
        'ACCEPTOR_SIGNING_DATE', 'acceptor_signing_date',
        'ACCEPTEE_SIGNING_DATE', 'acceptee_signing_date'
    ]

    def process_paragraph(paragraph):
        # We invoke this for every placeholder.
        for key, value in placeholders.items():
            placeholder = f"{{{{{key}}}}}"  # {{KEY}}

            # Quick check if placeholder exists in the full text at all
            if placeholder in paragraph.text:

                # Special handling for signatures (images)
                if key in signature_keys and value:
                    print(f"Found signature placeholder: {key}")
                    found_sig = False
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, "")
                            img_stream = fetch_image(value)
                            if img_stream:
                                run.add_picture(img_stream, width=Inches(1.5))
                                print(f"Signature inserted for {key}")
                                found_sig = True

                    if not found_sig:
                        # If signature placeholder is split across runs
                        print(
                            f"Signature placeholder {key} split across runs. clearing and appending.")
                        paragraph.text = paragraph.text.replace(
                            placeholder, "")
                        run = paragraph.add_run()
                        img_stream = fetch_image(value)
                        if img_stream:
                            run.add_picture(img_stream, width=Inches(1.5))

                elif key in highlighted_keys and value:
                    # Bold + Yellow highlight for acceptee fields (same as creator fields)
                    print(f"Applying bold + highlight to: {key}")
                    replaced_in_run = False
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            # Preserve original font properties
                            original_font_name = run.font.name
                            original_font_size = run.font.size

                            run.text = run.text.replace(
                                # Convert to uppercase like creator name
                                placeholder, str(value).upper())
                            run.bold = True
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            # Restore font properties to match original
                            if original_font_name:
                                run.font.name = original_font_name
                            if original_font_size:
                                run.font.size = original_font_size
                            replaced_in_run = True

                    if not replaced_in_run:
                        # Split across runs - get font from first run if available
                        original_font_name = None
                        original_font_size = None
                        if paragraph.runs:
                            first_run = paragraph.runs[0]
                            original_font_name = first_run.font.name
                            original_font_size = first_run.font.size

                        paragraph.text = paragraph.text.replace(
                            placeholder, "")
                        new_run = paragraph.add_run(
                            str(value).upper())  # Convert to uppercase
                        new_run.bold = True
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        # Apply original font properties
                        if original_font_name:
                            new_run.font.name = original_font_name
                        if original_font_size:
                            new_run.font.size = original_font_size
                else:
                    # Text replacement
                    replaced_in_run = False
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(
                                placeholder, str(value))
                            replaced_in_run = True

                    # If not found in any single run, it means it's split across runs
                    if not replaced_in_run:
                        paragraph.text = paragraph.text.replace(
                            placeholder, str(value))

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Save filled document
    filled_path = doc_path.replace('.docx', '_filled.docx')
    doc.save(filled_path)

    print(f"Filled template saved to: {filled_path}")
    return filled_path


def convert_to_pdf(docx_path: str) -> str:
    """
    Convert .docx to PDF using Microsoft Word.
    Returns path to the PDF file.
    """
    print(f"Converting to PDF...")

    pdf_path = docx_path.replace('.docx', '.pdf')

    try:
        convert(docx_path, pdf_path)
        print(f"PDF created: {pdf_path}")
        return pdf_path
    except AttributeError as e:
        # Known issue: Word.Application.Quit fails but PDF is created successfully
        if "Quit" in str(e) and os.path.exists(pdf_path):
            print(f"Word quit error ignored, PDF created: {pdf_path}")
            return pdf_path
        print(f"PDF conversion failed: {e}")
        raise
    except Exception as e:
        # Check if PDF was created despite error
        if os.path.exists(pdf_path):
            print(f"Error occurred but PDF exists: {pdf_path}")
            return pdf_path
        print(f"PDF conversion failed: {e}")
        raise


def upload_pdf(pdf_path: str, contract_id: str) -> str:
    """
    Upload generated PDF to Supabase Storage (contract_pdf bucket).
    Returns public URL.
    """
    supabase = get_supabase_client()

    # Storage path in the PDF bucket
    storage_path = f"{GENERATED_FOLDER}/{contract_id}.pdf"

    print(f"Uploading PDF to bucket '{PDF_BUCKET}': {storage_path}")

    # Read PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()

    # Upload to Supabase PDF bucket
    response = supabase.storage.from_(PDF_BUCKET).upload(
        storage_path,
        pdf_data,
        file_options={"content-type": "application/pdf"}
    )

    # Get public URL
    public_url = supabase.storage.from_(
        PDF_BUCKET).get_public_url(storage_path)

    print(f"PDF uploaded: {public_url}")
    return public_url


# Keys that should NOT be filled during initial contract creation
# These are the SIGNATURE SECTION fields - filled only when the creator/acceptor signs
# Note: Body placeholders like creator_name, acceptee_name ARE filled during creation
SIGNING_FIELDS_TO_EXCLUDE = [
    # Creator signature section fields (filled when creator signs)
    'creator_signature', 'CREATOR_SIGNATURE',
    'creator_signature_name', 'CREATOR_SIGNATURE_NAME',
    'creator_signature_id', 'CREATOR_SIGNATURE_ID',
    'creator_signature_date', 'CREATOR_SIGNATURE_DATE',
    'signing_date', 'SIGNING_DATE', 'creator_signing_date', 'CREATOR_SIGNING_DATE',
    # Acceptor/Acceptee signature section fields (filled when acceptee signs)
    'acceptor_signature', 'ACCEPTOR_SIGNATURE', 'acceptee_signature', 'ACCEPTEE_SIGNATURE',
    'acceptor_signature_name', 'ACCEPTOR_SIGNATURE_NAME', 'acceptee_signature_name', 'ACCEPTEE_SIGNATURE_NAME',
    'acceptor_signature_id', 'ACCEPTOR_SIGNATURE_ID', 'acceptee_signature_id', 'ACCEPTEE_SIGNATURE_ID',
    'acceptor_signature_date', 'ACCEPTOR_SIGNATURE_DATE', 'acceptee_signature_date', 'ACCEPTEE_SIGNATURE_DATE',
    'acceptor_signing_date', 'ACCEPTOR_SIGNING_DATE', 'acceptee_signing_date', 'ACCEPTEE_SIGNING_DATE',
]


def exclude_signing_fields(placeholders: dict) -> dict:
    """
    Remove signing-related fields from placeholders.
    These fields (creator name, IC, signature, date, and acceptor equivalents)
    should remain as placeholders in the document until actual signing.
    """
    return {
        key: value for key, value in placeholders.items()
        if key not in SIGNING_FIELDS_TO_EXCLUDE
    }


def generate_contract(template_name: str, placeholders: dict, contract_id: str) -> dict:
    """
    Full workflow: download template -> fill placeholders -> convert to PDF -> upload.
    Returns dict with success status and PDF URL.
    """
    try:
        # Step 0: Map placeholders
        mapping = get_template_mapping(template_name)
        print(f"\nDEBUG: Template ID: {template_name}")
        print(f"DEBUG: Mapping Config: {mapping}")
        print(f"DEBUG: Raw Placeholders (Keys): {list(placeholders.keys())}")

        mapped_placeholders = map_placeholders(placeholders, mapping)

        # Remove signing fields - these stay as placeholders until actual signing
        mapped_placeholders = exclude_signing_fields(mapped_placeholders)

        print(
            f"DEBUG: Mapped Placeholders (after excluding signing fields): {mapped_placeholders}")
        print(f"Mapped placeholders for {template_name}")

        # Step 1: Download template
        doc_path = download_template(template_name)

        # Step 2: Fill placeholders
        filled_path = fill_template(doc_path, mapped_placeholders)

        # Step 3: Convert to PDF
        pdf_path = convert_to_pdf(filled_path)

        # Step 4: Upload PDF
        pdf_url = upload_pdf(pdf_path, contract_id)

        # Cleanup temp files
        cleanup_temp_files([doc_path, filled_path, pdf_path])

        return {
            "success": True,
            "pdf_url": pdf_url,
            "contract_id": contract_id
        }

    except Exception as e:
        print(f"Contract generation failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }


def cleanup_temp_files(file_paths: list):
    """Remove temporary files"""
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass


def prepare_contract(template_name: str, placeholders: dict) -> dict:
    """
    Prepare contract by downloading template, filling placeholders, and converting to PDF.
    Stores PDF in prepared_contracts folder for later retrieval.
    Returns dict with prepare_id to retrieve the PDF later.
    """
    import uuid

    try:
        # Generate unique ID for this prepared contract
        prepare_id = str(uuid.uuid4())[:8]

        print(f"Preparing contract: {template_name} (ID: {prepare_id})")

        # Step 0: Map placeholders
        mapping = get_template_mapping(template_name)
        mapped_placeholders = map_placeholders(placeholders, mapping)

        # Remove signing fields - these stay as placeholders until actual signing
        mapped_placeholders = exclude_signing_fields(mapped_placeholders)

        # Step 1: Download template
        doc_path = download_template(template_name)

        # Step 2: Fill placeholders
        filled_path = fill_template(doc_path, mapped_placeholders)

        # Step 3: Convert to PDF
        pdf_path = convert_to_pdf(filled_path)

        # Step 4: Move PDF to prepared folder with prepare_id
        prepared_pdf_path = os.path.join(PREPARED_FOLDER, f"{prepare_id}.pdf")
        os.rename(pdf_path, prepared_pdf_path)

        # Cleanup temp files (but keep the prepared PDF)
        cleanup_temp_files([doc_path, filled_path])

        print(f"Contract prepared: {prepared_pdf_path}")

        return {
            "success": True,
            "prepare_id": prepare_id,
            "message": "Contract prepared successfully"
        }

    except Exception as e:
        print(f"Contract preparation failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }


def get_prepared_contract(prepare_id: str) -> str:
    """
    Get the path to a prepared contract PDF.
    Returns the file path or None if not found.
    """
    pdf_path = os.path.join(PREPARED_FOLDER, f"{prepare_id}.pdf")
    if os.path.exists(pdf_path):
        return pdf_path
    return None


def cleanup_prepared_contract(prepare_id: str):
    """Remove a prepared contract after it's been used"""
    pdf_path = os.path.join(PREPARED_FOLDER, f"{prepare_id}.pdf")
    if os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
            print(f"Cleaned up prepared contract: {prepare_id}")
        except Exception as e:
            print(f"Failed to cleanup prepared contract: {e}")


def save_signature_image(signature_base64: str, signature_id: str) -> str:
    """
    Save base64 signature image to temp folder.
    Returns the local file path.
    """
    try:
        # Remove data URL prefix if present
        if ',' in signature_base64:
            signature_base64 = signature_base64.split(',')[1]

        # Decode base64
        signature_data = base64.b64decode(signature_base64)

        # Save to file
        signature_path = os.path.join(SIGNATURES_FOLDER, f"{signature_id}.png")
        with open(signature_path, 'wb') as f:
            f.write(signature_data)

        print(f"Signature saved: {signature_path}")
        return signature_path
    except Exception as e:
        print(f"Failed to save signature: {e}")
        return None


def generate_signed_contract(
    template_name: str,
    placeholders: dict,
    creator_signature_base64: str,
    creator_name: str,
    creator_ic: str,
    contract_id: str
) -> str:
    """
    Generate a signed version of the contract with:
    - Creator signature image
    - Timestamp of signing
    - Creator name and IC
    Returns path to the signed PDF.
    """
    from datetime import datetime

    try:
        print(f"Generating signed contract: {contract_id}")

        # Save signature to temp file
        signature_path = save_signature_image(
            creator_signature_base64, f"sig_{contract_id}")

        # Get current timestamp with date and time (YYYY-MM-DD HH:MM:SS)
        signing_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Update placeholders with signature info
        signed_placeholders = {**placeholders}

        # Add creator signature (will be inserted as image)
        if signature_path and os.path.exists(signature_path):
            with open(signature_path, 'rb') as f:
                sig_data = f.read()
            sig_base64 = f"data:image/png;base64,{base64.b64encode(sig_data).decode()}"
            signed_placeholders['CREATOR_SIGNATURE'] = sig_base64
            signed_placeholders['creator_signature'] = sig_base64

        # Add signing details - support both uppercase and lowercase
        signed_placeholders['SIGNING_DATE'] = signing_timestamp
        signed_placeholders['signing_date'] = signing_timestamp
        signed_placeholders['CREATOR_SIGNING_DATE'] = signing_timestamp
        signed_placeholders['creator_signing_date'] = signing_timestamp

        # Creator signature section fields (new placeholders from template)
        signed_placeholders['creator_signature_name'] = creator_name
        signed_placeholders['CREATOR_SIGNATURE_NAME'] = creator_name
        signed_placeholders['creator_signature_id'] = creator_ic
        signed_placeholders['CREATOR_SIGNATURE_ID'] = creator_ic
        signed_placeholders['creator_signature_date'] = signing_timestamp
        signed_placeholders['CREATOR_SIGNATURE_DATE'] = signing_timestamp

        # Legacy creator name/IC fields (for backward compatibility)
        signed_placeholders['CREATOR_NAME'] = creator_name
        signed_placeholders['creator_name'] = creator_name
        signed_placeholders['CREATOR_IC'] = creator_ic
        signed_placeholders['creator_ic'] = creator_ic
        signed_placeholders['creator_id_number'] = creator_ic

        # Download template fresh
        doc_path = download_template(template_name)

        # Apply mapping to ensure frontend keys match Word template placeholders
        mapping = get_template_mapping(template_name)
        mapped_placeholders = map_placeholders(signed_placeholders, mapping)

        print(
            f"DEBUG: Signed contract placeholders: {list(mapped_placeholders.keys())}")

        # Fill with updated placeholders including signature
        filled_path = fill_template(doc_path, mapped_placeholders)

        # Convert to PDF
        pdf_path = convert_to_pdf(filled_path)

        # Cleanup temp files
        cleanup_temp_files([doc_path, filled_path])
        if signature_path and os.path.exists(signature_path):
            cleanup_temp_files([signature_path])

        print(f"Signed contract generated: {pdf_path}")
        return pdf_path

    except Exception as e:
        print(f"Failed to generate signed contract: {e}")
        import traceback
        traceback.print_exc()
        return None


# Preview function (returns filled docx without converting to PDF)
def preview_contract(template_name: str, placeholders: dict) -> str:
    """
    Generate preview of contract (filled .docx).
    Returns local path to filled document.
    """
    try:
        # Step 0: Map placeholders
        print(f"\nDEBUG: Previewing template: {template_name}")
        mapping = get_template_mapping(template_name)
        print(f"DEBUG: Preview Mapping Config: {mapping}")
        print(f"DEBUG: Raw Preview Placeholders: {placeholders}")

        mapped_placeholders = map_placeholders(placeholders, mapping)

        # Remove signing fields - these stay as placeholders until actual signing
        mapped_placeholders = exclude_signing_fields(mapped_placeholders)
        print(
            f"DEBUG: Mapped Preview Placeholders (after excluding signing fields): {mapped_placeholders}")

        doc_path = download_template(template_name)
        filled_path = fill_template(doc_path, mapped_placeholders)
        return filled_path
    except Exception as e:
        print(f"Preview failed: {e}")
        import traceback
        traceback.print_exc()
        raise


def extract_contract_text(template_name: str, placeholders: dict) -> dict:
    """
    Extract plain text from filled DOCX template.
    Returns dict with success status and extracted text.
    """
    try:
        print(f"📄 Extracting text from template: {template_name}")
        
        # Generate filled docx (same as preview)
        filled_path = preview_contract(template_name, placeholders)
        
        # Open the filled document and extract text
        doc = Document(filled_path)
        
        lines = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    lines.append(" | ".join(row_texts))
        
        full_text = "\n".join(lines)
        
        # Cleanup temp file
        doc_path = filled_path.replace('_filled.docx', '.docx')
        cleanup_temp_files([doc_path, filled_path])
        
        print(f"✅ Extracted {len(full_text)} characters from contract")
        
        return {
            "success": True,
            "text": full_text,
            "character_count": len(full_text)
        }
    
    except Exception as e:
        print(f"Text extraction failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "text": ""
        }


def upload_contract_pdf(pdf_path: str, user_id: str, contract_id: str) -> str:
    """
    Upload contract PDF to Supabase Storage under user_id folder.
    Storage path: {user_id}/{contract_id}.pdf
    Returns public URL.
    """
    supabase = get_supabase_client()

    # Storage path: user_id/contract_id.pdf
    storage_path = f"{user_id}/{contract_id}.pdf"

    print(f"Uploading PDF to bucket '{PDF_BUCKET}': {storage_path}")

    # Read PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()

    # Upload to Supabase PDF bucket
    try:
        response = supabase.storage.from_(PDF_BUCKET).upload(
            storage_path,
            pdf_data,
            file_options={"content-type": "application/pdf"}
        )
    except Exception as e:
        # If file already exists, try to update it
        if "already exists" in str(e).lower() or "duplicate" in str(e).lower():
            print(f"File exists, updating: {storage_path}")
            response = supabase.storage.from_(PDF_BUCKET).update(
                storage_path,
                pdf_data,
                file_options={"content-type": "application/pdf"}
            )
        else:
            raise e

    # Get public URL
    public_url = supabase.storage.from_(
        PDF_BUCKET).get_public_url(storage_path)

    print(f"PDF uploaded: {public_url}")
    return public_url


def create_contract_record(contract_data: dict) -> dict:
    """
    Create a new contract record in Supabase contracts table.
    Returns the created contract data.
    """
    supabase = get_supabase_client()

    print(f"Creating contract record: {contract_data.get('contract_id')}")

    # Insert into contracts table
    result = supabase.table('contracts').insert(contract_data).execute()

    if result.data:
        print(f"Contract record created: {result.data[0].get('contract_id')}")
        return result.data[0]
    else:
        raise Exception("Failed to create contract record")


def finalize_contract(
    prepare_id: str,
    user_id: str,
    acceptee_id: str,
    contract_name: str,
    contract_topic: str,
    template_type: str,
    form_data: dict,
    creator_signature: str = None,
    creator_name: str = None,
    creator_ic: str = None,
    creator_nfc_verified: bool = False,
    creator_face_verified: bool = False,
    due_date: str = None
) -> dict:
    """
    Finalize contract creation:
    1. Generate signed PDF with creator signature, timestamp, and details
    2. Upload PDF to storage under user_id folder
    3. Create contract record in database
    4. Cleanup prepared contract
    Returns dict with contract data and PDF URL.
    """
    import uuid
    from datetime import datetime, timedelta

    try:
        # Generate contract ID
        contract_id = f"CNT-{datetime.now().strftime('%Y%m%d%H%M%S')}-{str(uuid.uuid4())[:4].upper()}"

        print(f"Finalizing contract: {contract_id}")

        # Step 1: Generate signed contract with signature embedded
        if creator_signature:
            print("Generating signed version with creator signature...")

            # Build placeholders from form_data
            # Use lowercase keys to match Word template placeholders
            placeholders = {**form_data}
            placeholders['creator_name'] = creator_name or ''
            placeholders['creator_id_number'] = creator_ic or ''

            # Generate signed PDF
            pdf_path = generate_signed_contract(
                template_name=template_type,
                placeholders=placeholders,
                creator_signature_base64=creator_signature,
                creator_name=creator_name or 'Unknown',
                creator_ic=creator_ic or 'Unknown',
                contract_id=contract_id
            )

            if not pdf_path:
                print("Signed contract generation failed, using prepared contract")
                pdf_path = get_prepared_contract(prepare_id)
        else:
            # No signature provided, use prepared contract
            pdf_path = get_prepared_contract(prepare_id)

        if not pdf_path:
            return {
                "success": False,
                "error": f"Contract PDF not found"
            }

        # Step 2: Upload PDF to storage under user_id folder
        pdf_url = upload_contract_pdf(pdf_path, user_id, contract_id)

        # Step 3: Create contract record in database
        # IMPORTANT: Save creator signature and signature section fields to form_data
        # so they can be used when acceptor signs
        from datetime import datetime as dt
        creator_signing_timestamp = dt.now().strftime("%Y-%m-%d %H:%M:%S")

        form_data_with_signature = {**form_data}
        if creator_signature:
            form_data_with_signature['creator_signature'] = creator_signature
        form_data_with_signature['creator_name'] = creator_name or ''
        form_data_with_signature['creator_id_number'] = creator_ic or ''
        # Save signature section fields for when acceptor signs and PDF is regenerated
        form_data_with_signature['creator_signature_name'] = creator_name or ''
        form_data_with_signature['creator_signature_id'] = creator_ic or ''
        form_data_with_signature['creator_signature_date'] = creator_signing_timestamp

        contract_data = {
            "contract_id": contract_id,
            "created_user_id": user_id,
            "acceptee_user_id": acceptee_id,
            "contract_name": contract_name,
            "contract_topic": contract_topic,
            "status": "Pending",
            "template_type": template_type,
            "form_data": form_data_with_signature,
            "pdf_url": pdf_url,
            "creator_nfc_verified": creator_nfc_verified,
            "creator_face_verified": creator_face_verified,
            "acceptee_nfc_verified": False,
            "acceptee_face_verified": False,
            "created_at": datetime.now().isoformat(),
            "due_date": due_date or (datetime.now() + timedelta(days=30)).isoformat(),
        }

        created_contract = create_contract_record(contract_data)

        # Step 4: Cleanup prepared contract and temp PDF
        cleanup_prepared_contract(prepare_id)
        if pdf_path and os.path.exists(pdf_path) and TEMP_FOLDER in pdf_path:
            cleanup_temp_files([pdf_path])

        return {
            "success": True,
            "contract": created_contract,
            "pdf_url": pdf_url,
            "contract_id": contract_id
        }

    except Exception as e:
        print(f"Contract finalization failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }


def update_contract_pdf(pdf_path: str, user_id: str, contract_id: str) -> str:
    """
    Update (overwrite) existing contract PDF in Supabase Storage.
    Storage path: {user_id}/{contract_id}.pdf
    Returns public URL.
    """
    supabase = get_supabase_client()

    # Storage path: user_id/contract_id.pdf
    storage_path = f"{user_id}/{contract_id}.pdf"

    print(f"📤 Updating PDF in bucket '{PDF_BUCKET}': {storage_path}")

    # Read PDF file
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()

    # Use update to overwrite existing file
    try:
        response = supabase.storage.from_(PDF_BUCKET).update(
            storage_path,
            pdf_data,
            file_options={"content-type": "application/pdf"}
        )
    except Exception as e:
        # If update fails, try upload (file might not exist)
        print(f"⚠️ Update failed, trying upload: {e}")
        response = supabase.storage.from_(PDF_BUCKET).upload(
            storage_path,
            pdf_data,
            file_options={"content-type": "application/pdf", "upsert": "true"}
        )

    # Get public URL with cache buster
    import time
    public_url = supabase.storage.from_(
        PDF_BUCKET).get_public_url(storage_path)
    # Add cache buster
    public_url_with_cache = f"{public_url}?t={int(time.time())}"

    print(f"✅ PDF updated: {public_url_with_cache}")
    return public_url_with_cache


def update_contract_record(contract_id: str, updates: dict) -> dict:
    """
    Update an existing contract record in Supabase contracts table.
    Returns the updated contract data.
    """
    supabase = get_supabase_client()

    print(f"📝 Updating contract record: {contract_id}")

    # Update in contracts table
    result = supabase.table('contracts').update(
        updates).eq('contract_id', contract_id).execute()

    if result.data:
        print(f"✅ Contract record updated: {contract_id}")
        return result.data[0]
    else:
        raise Exception("Failed to update contract record")


def get_contract_by_id(contract_id: str) -> dict:
    """
    Fetch a contract record from Supabase by contract_id.
    Returns the contract data or None.
    """
    supabase = get_supabase_client()

    result = supabase.table('contracts').select(
        '*').eq('contract_id', contract_id).execute()

    if result.data and len(result.data) > 0:
        return result.data[0]
    return None


def sign_contract_acceptor(
    contract_id: str,
    acceptor_signature_base64: str,
    acceptor_name: str,
    acceptor_ic: str,
    acceptor_nfc_verified: bool = False,
    acceptor_face_verified: bool = False
) -> dict:
    """
    Sign contract as acceptor:
    1. Fetch contract record from database
    2. Re-download the template
    3. Fill in ALL placeholders including:
       - Creator details (from form_data)
       - Acceptor details (from this call)
       - Both signatures
    4. Generate new PDF with both signatures
    5. Overwrite the existing PDF in storage
    6. Update contract record with acceptor signature info and status='Ongoing'
    Returns dict with success status and updated PDF URL.
    """
    from datetime import datetime

    try:
        print(f"✍️ Acceptor signing contract: {contract_id}")

        # Step 1: Fetch existing contract
        contract = get_contract_by_id(contract_id)
        if not contract:
            return {
                "success": False,
                "error": f"Contract not found: {contract_id}"
            }

        template_type = contract.get('template_type')
        form_data = contract.get('form_data', {})
        creator_id = contract.get('created_user_id')
        pdf_url = contract.get('pdf_url', '')

        print(f"📋 Contract template: {template_type}")

        # Save acceptor signature to temp file
        acceptor_sig_path = save_signature_image(
            acceptor_signature_base64, f"acceptor_sig_{contract_id}")

        # Get current timestamp with date and time (YYYY-MM-DD HH:MM:SS)
        signing_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Step 3: Build complete placeholders with all details
        placeholders = {**form_data}

        # IMPORTANT: Preserve creator signature from form_data if it exists
        # The creator signature should have been saved during contract creation
        creator_sig = form_data.get(
            'creator_signature') or form_data.get('CREATOR_SIGNATURE')
        if creator_sig:
            print(f"📝 Preserving creator signature from form_data")
            placeholders['CREATOR_SIGNATURE'] = creator_sig
            placeholders['creator_signature'] = creator_sig

        # Preserve creator signature section fields (name, id, date)
        creator_sig_name = form_data.get(
            'creator_signature_name') or form_data.get('CREATOR_SIGNATURE_NAME')
        creator_sig_id = form_data.get(
            'creator_signature_id') or form_data.get('CREATOR_SIGNATURE_ID')
        creator_sig_date = form_data.get(
            'creator_signature_date') or form_data.get('CREATOR_SIGNATURE_DATE')

        if creator_sig_name:
            print(f"📝 Preserving creator signature name: {creator_sig_name}")
            placeholders['creator_signature_name'] = creator_sig_name
            placeholders['CREATOR_SIGNATURE_NAME'] = creator_sig_name
        if creator_sig_id:
            print(f"📝 Preserving creator signature id: {creator_sig_id}")
            placeholders['creator_signature_id'] = creator_sig_id
            placeholders['CREATOR_SIGNATURE_ID'] = creator_sig_id
        if creator_sig_date:
            print(f"📝 Preserving creator signature date: {creator_sig_date}")
            placeholders['creator_signature_date'] = creator_sig_date
            placeholders['CREATOR_SIGNATURE_DATE'] = creator_sig_date

        # Add acceptor details - support both uppercase and lowercase
        placeholders['ACCEPTOR_NAME'] = acceptor_name
        placeholders['acceptor_name'] = acceptor_name
        placeholders['ACCEPTEE_NAME'] = acceptor_name
        placeholders['acceptee_name'] = acceptor_name

        placeholders['ACCEPTOR_IC'] = acceptor_ic
        placeholders['acceptor_ic'] = acceptor_ic
        placeholders['ACCEPTEE_IC'] = acceptor_ic
        placeholders['acceptee_ic'] = acceptor_ic
        placeholders['acceptor_id_number'] = acceptor_ic
        placeholders['acceptee_id_number'] = acceptor_ic

        # Add acceptor signature
        if acceptor_sig_path and os.path.exists(acceptor_sig_path):
            with open(acceptor_sig_path, 'rb') as f:
                sig_data = f.read()
            sig_base64 = f"data:image/png;base64,{base64.b64encode(sig_data).decode()}"
            placeholders['ACCEPTOR_SIGNATURE'] = sig_base64
            placeholders['acceptor_signature'] = sig_base64
            placeholders['ACCEPTEE_SIGNATURE'] = sig_base64
            placeholders['acceptee_signature'] = sig_base64

        # Add acceptor signing date
        placeholders['ACCEPTOR_SIGNING_DATE'] = signing_timestamp
        placeholders['acceptor_signing_date'] = signing_timestamp
        placeholders['ACCEPTEE_SIGNING_DATE'] = signing_timestamp
        placeholders['acceptee_signing_date'] = signing_timestamp

        # Acceptee signature section fields (new placeholders from template)
        placeholders['acceptee_signature_name'] = acceptor_name
        placeholders['ACCEPTEE_SIGNATURE_NAME'] = acceptor_name
        placeholders['acceptor_signature_name'] = acceptor_name
        placeholders['ACCEPTOR_SIGNATURE_NAME'] = acceptor_name
        placeholders['acceptee_signature_id'] = acceptor_ic
        placeholders['ACCEPTEE_SIGNATURE_ID'] = acceptor_ic
        placeholders['acceptor_signature_id'] = acceptor_ic
        placeholders['ACCEPTOR_SIGNATURE_ID'] = acceptor_ic
        placeholders['acceptee_signature_date'] = signing_timestamp
        placeholders['ACCEPTEE_SIGNATURE_DATE'] = signing_timestamp
        placeholders['acceptor_signature_date'] = signing_timestamp
        placeholders['ACCEPTOR_SIGNATURE_DATE'] = signing_timestamp

        print(f"📝 Filled {len(placeholders)} placeholders")

        # Step 4: Download template and generate new PDF
        doc_path = download_template(template_type)
        filled_path = fill_template(doc_path, placeholders)
        pdf_path = convert_to_pdf(filled_path)

        # Step 5: Upload/overwrite PDF in storage
        new_pdf_url = update_contract_pdf(pdf_path, creator_id, contract_id)

        # Step 6: Update contract record
        # Note: acceptee_signed_at column doesn't exist in current schema
        updates = {
            "status": "Ongoing",
            "acceptee_nfc_verified": acceptor_nfc_verified,
            "acceptee_face_verified": acceptor_face_verified,
            "pdf_url": new_pdf_url,
        }

        updated_contract = update_contract_record(contract_id, updates)

        # Cleanup temp files
        cleanup_temp_files([doc_path, filled_path, pdf_path])
        if acceptor_sig_path:
            cleanup_temp_files([acceptor_sig_path])

        print(f"✅ Contract signed by acceptor: {contract_id}")

        return {
            "success": True,
            "contract": updated_contract,
            "pdf_url": new_pdf_url,
            "message": "Contract signed successfully"
        }

    except Exception as e:
        print(f"❌ Acceptor signing failed: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e)
        }
